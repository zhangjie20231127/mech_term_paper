from __future__ import annotations

from pathlib import Path
from typing import List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "paper_submitable_source.md"
OUTPUT = ROOT / "期末论文.docx"


def clear_document(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_run_font(run, east_asia: str, ascii_font: str, size_pt: int, bold: bool = False) -> None:
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size_pt)
    run.bold = bold


def add_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.replace("`", ""))
    set_run_font(run, "宋体", "Times New Roman", 12, bold=bold)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.replace("`", ""))
    if level == 1:
        set_run_font(run, "黑体", "Arial", 14, bold=True)
    else:
        set_run_font(run, "黑体", "Arial", 12, bold=True)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text.replace("`", ""))
    set_run_font(run, "黑体", "Arial", 16, bold=True)


def add_keywords(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.replace("`", ""))
    set_run_font(run, "宋体", "Times New Roman", 12, bold=True)


def parse_table(lines: List[str]) -> List[List[str]]:
    rows: List[List[str]] = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if all(cell.replace("-", "").replace(":", "") == "" for cell in cells):
            continue
        rows.append(cells)
    return rows


def add_table(doc: Document, rows: List[List[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text.replace("`", ""))
            set_run_font(run, "宋体", "Times New Roman", 10, bold=(r_idx == 0))
    doc.add_paragraph()


def build_doc() -> None:
    source_lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    clear_document(doc)

    i = 0
    while i < len(source_lines):
        line = source_lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("# "):
            add_title(doc, stripped[2:].strip())
            i += 1
            continue
        if stripped.startswith("## "):
            heading = stripped[3:].strip()
            if heading == "摘要":
                add_heading(doc, heading, 1)
            elif heading == "参考文献":
                add_heading(doc, heading, 1)
            else:
                add_heading(doc, heading, 1)
            i += 1
            continue
        if stripped.startswith("### "):
            add_heading(doc, stripped[4:].strip(), 2)
            i += 1
            continue
        if stripped.startswith("|"):
            table_lines = []
            while i < len(source_lines) and source_lines[i].strip().startswith("|"):
                table_lines.append(source_lines[i])
                i += 1
            add_table(doc, parse_table(table_lines))
            continue
        if stripped.startswith("**关键词：**"):
            add_keywords(doc, stripped.replace("**", ""))
            i += 1
            continue

        add_paragraph(doc, stripped)
        i += 1

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
